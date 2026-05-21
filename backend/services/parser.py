from __future__ import annotations

import io
import logging
import re
import unicodedata
from dataclasses import dataclass

import pypdf
import pdfplumber
import docx

log = logging.getLogger(__name__)


_PREVIEW_CHARS = 500

_PDF_TYPES  = {"application/pdf"}
_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_DOC_TYPES  = {"application/msword"}


@dataclass(frozen=True, slots=True)
class ParsedResume:
    filename: str
    size_bytes: int
    detected_type: str
    full_text: str
    raw_text_for_contacts: str
    extracted_text_preview: str
    used_fallback_parser: bool = False  # True when pdfplumber was used instead of pypdf


def parse_resume(contents: bytes, filename: str, content_type: str) -> ParsedResume:
    """
    Dispatch to the correct parser based on content_type.
    Extracts raw text AND a pre-clean snapshot for contact extraction.
    """
    used_fallback = False
    raw_text = ""
    if content_type in _PDF_TYPES:
        text, used_fallback, raw_text = _parse_pdf(contents)
        detected = "pdf"
    elif content_type in _DOCX_TYPES:
        text = _parse_docx(contents)
        raw_text = text
        detected = "docx"
    elif content_type in _DOC_TYPES:
        text = contents.decode("latin-1", errors="replace")
        raw_text = text
        detected = "doc"
    else:
        text = ""
        raw_text = ""
        detected = "unknown"

    return ParsedResume(
        filename=filename,
        size_bytes=len(contents),
        detected_type=detected,
        full_text=text,
        raw_text_for_contacts=raw_text,
        extracted_text_preview=text[:_PREVIEW_CHARS].strip(),
        used_fallback_parser=used_fallback,
    )


# ---------------------------------------------------------------------------
# Internal parsers
# ---------------------------------------------------------------------------

# Minimum number of recognised standalone headings required before
# the fallback parser is skipped.  Resumes with fewer than this count
# are considered low-confidence and trigger pdfplumber.
_HEADING_CONFIDENCE_THRESHOLD = 2

# Standalone heading detector used only for confidence scoring.
# Intentionally narrow: the five highest-signal ATS section names.
_SCORE_HEADING_PATTERN = re.compile(
    r"(?im)^[ \t]*(?:skills?|experience|projects?|education|certifications?)[ \t]*:?[ \t]*$"
)


def _score_headings(text: str) -> int:
    """Count distinct standalone major headings in *text*."""
    return len(_SCORE_HEADING_PATTERN.findall(text))


def _extract_pypdf(contents: bytes) -> str:
    """Extract raw text from all pages using pypdf."""
    reader = pypdf.PdfReader(io.BytesIO(contents))
    pages: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            pages.append(extracted)
    return "\n".join(pages)


def _extract_pdfplumber(contents: bytes) -> str:
    """Extract raw text from all pages using pdfplumber."""
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                pages.append(extracted)
    return "\n".join(pages)


def _parse_pdf(contents: bytes) -> tuple[str, bool, str]:
    """
    Returns (cleaned_text, used_fallback_parser, raw_text_for_contacts).
    Tries pypdf first. Falls back to pdfplumber for multi-column/graphic resumes.
    No external API calls.
    """
    pypdf_raw     = _extract_pypdf(contents)
    pypdf_cleaned = _clean_pdf_text(pypdf_raw)
    pypdf_score   = _score_headings(pypdf_cleaned)

    if pypdf_score >= _HEADING_CONFIDENCE_THRESHOLD:
        return pypdf_cleaned, False, pypdf_raw

    plumber_raw     = _extract_pdfplumber(contents)
    plumber_cleaned = _clean_pdf_text(plumber_raw)
    plumber_score   = _score_headings(plumber_cleaned)

    if plumber_score >= pypdf_score:
        log.info("parser: pdfplumber gave better result (%d vs %d headings)", plumber_score, pypdf_score)
        return plumber_cleaned, True, plumber_raw

    # pypdf was still better even though below threshold — use it
    return pypdf_cleaned, False, pypdf_raw


def _parse_docx(contents: bytes) -> str:
    doc = docx.Document(io.BytesIO(contents))
    parts: list[str] = []
    # headers/footers — contact info often lives here
    for section in doc.sections:
        for container in [section.header, section.footer]:
            for p in container.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
    # body paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # tables — skills grids and experience tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# PDF text cleanup pipeline
# ---------------------------------------------------------------------------

# Pass 3 — camelCase boundary: lowercase immediately followed by uppercase
# e.g. "PythonDeveloper" → "Python Developer"
_CAMEL_SPLIT = re.compile(r"([a-z])([A-Z])")

# Pass 3 — space before opening paren when preceded by a word char
# e.g. "Python(3.x)" → "Python (3.x)"
_SPACE_BEFORE_PAREN = re.compile(r"(\w)\(")

# Pass 3 — space after closing paren when followed by a word char
# e.g. "(remote)Role" → "(remote) Role"
_SPACE_AFTER_PAREN = re.compile(r"\)(\w)")

# Pass 3 — space after comma/semicolon when not already followed by whitespace
_SPACE_AFTER_COMMA = re.compile(r"([,;])(\S)")

# Pass 3 — space around forward-slash between word characters
# e.g. "React/Node" → "React / Node"  (common in skill lists)
_SPACE_AROUND_SLASH = re.compile(r"(\w)/(\w)")

# Pass 2 — decorative / box-drawing / dingbat codepoint ranges
# Covers: Box Drawing U+2500-257F, Block Elements U+2580-259F,
#         Geometric Shapes U+25A0-25FF (except ◦ U+25E6 kept as bullet),
#         Dingbats U+2700-27BF, Ornamental Dashes U+2E00-2E7F,
#         Enclosed Alphanumerics U+2460-24FF
_DECORATIVE = re.compile(
    r"[\u2300-\u23ff"   # miscellaneous technical (⌢ and similar)
    r"\u2400-\u243f"   # control pictures
    r"\u2440-\u245f"   # optical character recognition
    r"\u2460-\u24ff"   # enclosed alphanumerics
    r"\u2500-\u257f"   # box drawing
    r"\u2580-\u259f"   # block elements
    r"\u25a0-\u25e5"   # geometric shapes (solid — keep ◦ U+25E6)
    r"\u25e7-\u25ff"   # rest of geometric shapes
    r"\u2600-\u26ff"   # miscellaneous symbols (☎ ✉ ♂ ♀ etc.)
    r"\u2700-\u27bf"   # dingbats
    r"\u2e00-\u2e7f"   # supplemental punctuation / ornamental
    r"\uf000-\ufffd"   # private use area (icon fonts: FontAwesome, etc.)
    r"]+"
)

# Pass 5 — known ATS section heading words that should be on their own line.
# Matches the word when it appears after non-newline content on the same line,
# preceded by 2+ spaces or a pipe/bullet separator.
# Only fires when the heading word is followed by end-of-line or a colon.
_HEADING_WORDS = (
    r"Skills?|Technical\s+Skills?|Core\s+Competenc(?:y|ies)"
    r"|Experience|Work\s+Experience|Professional\s+Experience"
    r"|Employment(?:\s+History)?|Internships?"
    r"|Education|Academic(?:\s+Background)?"
    r"|Projects?|Personal\s+Projects?|Academic\s+Projects?"
    r"|Certifications?|Courses?|Awards?"
    r"|Summary|Objective|Profile"
    r"|Publications?|Volunteer|References?"
    r"|Languages?|Tools?(?:\s+(?:&|and)\s+Technologies?)?"
    r"|Tech(?:nology)?\s+Stack|Platforms?"
)
_HEADING_INLINE = re.compile(
    rf"(?<=[^\n])[ \t]*[|•·–—]?[ \t]+({_HEADING_WORDS})[ \t]*:?[ \t]*(?=\n|$)",
    re.IGNORECASE,
)


def _clean_pdf_text(raw: str) -> str:
    """
    Six-pass cleanup pipeline for raw pypdf output.

    Pass order is intentional — each pass assumes the previous one has run.
    DOCX output bypasses this entirely.
    """
    text = raw

    # ------------------------------------------------------------------
    # Pass 1 — Unicode normalisation + bullet/dash/separator mapping
    # ------------------------------------------------------------------
    text = unicodedata.normalize("NFC", text)

    # Map unicode bullets to ASCII hyphen so heading patterns see plain text
    text = text.translate(_UNICODE_BULLET_TABLE)

    # ------------------------------------------------------------------
    # Pass 2 — Strip decorative symbols (box-drawing, dingbats, etc.)
    #          Replace with a space so adjacent words don't merge
    # ------------------------------------------------------------------
    text = _DECORATIVE.sub(" ", text)

    # Strip icon-font slash prefixes left after glyph removal
    # e.g. "-/ envelpe" → removes the "-/" artifact from contact headers
    text = re.sub(r"(?<=[\s])-/\s*", " ", text)

    # Pass 3 — Restore missing spaces
    # Guard: protect emails AND phones before applying space fixes
    _email_placeholders: list[str] = []
    _phone_placeholders: list[str] = []

    def _mask_contacts(t: str) -> str:
        import re as _re
        _simple_email = _re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
        _simple_phone = _re.compile(r"\+?\d[\d\s\-().]{8,18}\d")
        for i, m in enumerate(_simple_email.finditer(t)):
            ph = f"__EMAIL{i}__"
            _email_placeholders.append(m.group(0))
            t = t.replace(m.group(0), ph, 1)
        for i, m in enumerate(_simple_phone.finditer(t)):
            ph = f"__PHONE{i}__"
            _phone_placeholders.append(m.group(0))
            t = t.replace(m.group(0), ph, 1)
        return t

    def _unmask_contacts(t: str) -> str:
        for i, val in enumerate(_email_placeholders):
            t = t.replace(f"__EMAIL{i}__", val)
        for i, val in enumerate(_phone_placeholders):
            t = t.replace(f"__PHONE{i}__", val)
        return t

    text = _mask_contacts(text)
    text = _CAMEL_SPLIT.sub(r"\1 \2", text)
    text = _SPACE_BEFORE_PAREN.sub(r"\1 (", text)
    text = _SPACE_AFTER_PAREN.sub(r") \1", text)
    text = _SPACE_AFTER_COMMA.sub(r"\1 \2", text)
    text = _SPACE_AROUND_SLASH.sub(r"\1 / \2", text)
    text = _unmask_contacts(text)

    # ------------------------------------------------------------------
    # Pass 3b — Rejoin ALL-CAPS words broken by PDF kerning/ligature
    #
    # e.g. "RELEV ANT" → "RELEVANT", "COURS EWORK" → "COURSEWORK"
    # Fires when both fragments are ALL-CAPS alpha and combined length ≤ 20.
    # ------------------------------------------------------------------
    text = re.sub(
        r"\b([A-Z]{2,})[ \t]+([A-Z]{2,})\b",
        lambda m: (
            m.group(1) + m.group(2)
            if len(m.group(1)) + len(m.group(2)) <= 20
            else m.group(0)
        ),
        text,
    )

    # ------------------------------------------------------------------
    # Pass 4 — Repair broken words split across lines by PDF layout
    #
    # Heuristic: if a line ends with a lowercase letter (no punctuation)
    # and the next non-empty line starts with a lowercase letter,
    # the line break is a layout artifact — rejoin with a space.
    # Guard: do NOT rejoin if the next line looks like a new sentence
    # (starts with uppercase) or is a standalone short token (heading).
    # ------------------------------------------------------------------
    text = re.sub(
        r"([a-z,])\n([a-z])",
        r"\1 \2",
        text,
    )

    # ------------------------------------------------------------------
    # Pass 5 — Isolate section headings onto their own lines
    #
    # When a known heading word appears mid-line after content,
    # inject a newline before it so the section slicer can anchor on it.
    # ------------------------------------------------------------------
    text = _HEADING_INLINE.sub(r"\n\1", text)

    # ------------------------------------------------------------------
    # Pass 6 — Whitespace collapse
    # ------------------------------------------------------------------
    # Strip trailing spaces/tabs from every line
    text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)
    # Collapse runs of spaces/tabs within a line to a single space
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Collapse 3+ consecutive blank lines to exactly 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip leading/trailing whitespace from the whole document
    text = text.strip()

    return text


# ---------------------------------------------------------------------------
# Unicode → ASCII translation table  (built once at import time)
# ---------------------------------------------------------------------------

def _build_bullet_table() -> dict[int, str]:
    """
    Map common unicode bullets, dashes, and vertical separators to their
    nearest ASCII equivalents so downstream regex patterns stay simple.
    """
    mapping: dict[int, str] = {}

    # Bullets → hyphen-minus
    for cp in [
        0x2022,  # • BULLET
        0x2023,  # ‣ TRIANGULAR BULLET
        0x25E6,  # ◦ WHITE BULLET
        0x2043,  # ⁃ HYPHEN BULLET
        0x2219,  # ∙ BULLET OPERATOR
        0x00B7,  # · MIDDLE DOT
        0x2027,  # ‧ HYPHENATION POINT
    ]:
        mapping[cp] = "-"

    # Dashes → hyphen-minus
    for cp in [
        0x2013,  # – EN DASH
        0x2014,  # — EM DASH
        0x2015,  # ― HORIZONTAL BAR
        0x2212,  # − MINUS SIGN
        0xFE58,  # ﹘ SMALL EM DASH
        0xFE63,  # ﹣ SMALL HYPHEN-MINUS
        0xFF0D,  # － FULLWIDTH HYPHEN-MINUS
    ]:
        mapping[cp] = "-"

    # Vertical separators → pipe
    for cp in [
        0x2502,  # │ BOX DRAWINGS LIGHT VERTICAL
        0x2503,  # ┃ BOX DRAWINGS HEAVY VERTICAL
        0xFF5C,  # ｜ FULLWIDTH VERTICAL LINE
        0x01C0,  # ǀ LATIN LETTER DENTAL CLICK
    ]:
        mapping[cp] = "|"

    # Non-breaking / zero-width spaces → regular space
    for cp in [
        0x00A0,  # NO-BREAK SPACE
        0x202F,  # NARROW NO-BREAK SPACE
        0x200B,  # ZERO WIDTH SPACE
        0xFEFF,  # ZERO WIDTH NO-BREAK SPACE (BOM)
        0x2009,  # THIN SPACE
        0x200A,  # HAIR SPACE
    ]:
        mapping[cp] = " "

    # Line/paragraph separators → newline
    mapping[0x2028] = "\n"  # LINE SEPARATOR
    mapping[0x2029] = "\n"  # PARAGRAPH SEPARATOR

    return mapping


_UNICODE_BULLET_TABLE: dict[int, str] = _build_bullet_table()
